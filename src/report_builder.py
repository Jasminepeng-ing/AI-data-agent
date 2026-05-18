"""
src/report_builder.py
=====================
报告生成模块（纯 Python，不依赖 Streamlit）。

提供两种报告格式：
- Word (.docx): python-docx 格式化，适合二次编辑
- PDF (.pdf):   weasyprint HTML→PDF 管线，含 Plotly 图表截图（需 kaleido）

对外入口：
  markdown_to_word(markdown_content, title) → bytes
  markdown_to_pdf(markdown_content, title, chart_figures, chart_captions) → bytes
"""

import re
import base64
from io import BytesIO
from datetime import datetime

# Emoji Unicode 区间：Microsoft YaHei / SimHei 等 CJK 字体无法渲染，会显示为方框
# Markdown 图片占位符：LLM 有时输出 ![图表N:caption]()，需识别并注入真实图表
_MD_IMG_RE = re.compile(r'^!\[([^\]]*)\]\([^)]*\)')

# 图表标签行（与函数内局部变量同模式，用于 _is_block_start）
_BLOCK_CHART_RE = re.compile(r'^图表?\s*\d+\s*[：:]')


def _is_block_start(line: str) -> bool:
    """判断一行是否为新块的起始（应中止软换行续行的累积）。"""
    s = line.strip()
    return (not s
            or bool(re.match(r'^#{1,6}[\s#]', s))
            or bool(re.match(r'^[-*] ', s))
            or bool(re.match(r'^\d+\. ', s))
            or s.startswith('|')
            or bool(re.match(r'^-{3,}\s*$', s))
            or s.startswith('>')
            or bool(_BLOCK_CHART_RE.match(s))
            or bool(_MD_IMG_RE.match(s)))

_EMOJI_STRIP_RE = re.compile(
    '[' + '\U0001F000-\U0001FFFF\U00002600-\U000027BF'
    + chr(0xFE0F)   # Variation Selector-16
    + chr(0x200D)   # Zero-Width Joiner
    + ']+'
)


# ── 函数 1: markdown_to_word ─────────────────────────────────────────────────

def markdown_to_word(
    markdown_content: str,
    title: str,
    chart_captions: list = None,
    chart_png_bytes: list = None,
) -> bytes:
    """
    把 Markdown 字符串转成格式化的 Word 文档，返回 bytes。

    支持：一/二/三级标题、正文、加粗、项目符号、Markdown 表格、水平线、图表嵌入
    chart_captions / chart_png_bytes 均传入时在报告末尾追加图表截图。
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── 页面设置 ──────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # ── 页眉：报告标题右对齐 ──────────────────────────────────────────────────
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = title
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.style.font.size = Pt(9)
    header_para.style.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 页脚：日期左 + 页码右 ─────────────────────────────────────────────────
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    # 左侧日期
    run_date = footer_para.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    run_date.font.size = Pt(9)
    run_date.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    # Tab 分隔到右侧
    footer_para.add_run("\t\t")
    # 右侧页码（Word 域）
    run_pg = footer_para.add_run()
    run_pg.font.size = Pt(9)
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run_pg._r.append(fldChar1); run_pg._r.append(instrText); run_pg._r.append(fldChar2)

    # ── 辅助：段落格式 ────────────────────────────────────────────────────────
    def _para(text: str, bold=False, size=11, space_before=0, space_after=4,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, color: tuple = None):
        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = Pt(size * 1.3)
        # 内联 **加粗** 解析
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            m = re.match(r'\*\*([^*]+)\*\*', part)
            run = p.add_run(m.group(1) if m else part)
            run.bold = bold or bool(m)
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def _add_hr(doc):
        """添加水平分割线段落。"""
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── 文档标题 ──────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(18)
    title_para.paragraph_format.space_after  = Pt(12)
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(24)

    # 生成日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(12)
    date_run = date_para.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    _add_hr(doc)

    # ── 按行解析 Markdown ─────────────────────────────────────────────────────
    lines = markdown_content.splitlines()
    i = 0
    _WORD_CHART_RE = re.compile(r'^图表?\s*\d+\s*[：:]')
    _word_unplaced = list(zip(chart_captions or [], chart_png_bytes or []))
    _word_chart_idx = [0]   # 当前待放置图表下标（顺序兜底）

    def _word_inject_chart(caption_hint: str):
        """在 Word 当前位置嵌入一张与 caption_hint 最匹配的图表。"""
        if not _word_unplaced:
            return
        # 简单匹配：找 caption 与 hint 共同字符数最多的
        best_idx, best_score = 0, -1
        for wi, (cap, png) in enumerate(_word_unplaced):
            if png is None:
                continue
            score = sum(1 for ch in caption_hint if ch in cap)
            if score > best_score:
                best_score, best_idx = score, wi
        cap, png = _word_unplaced.pop(best_idx)
        if png is None:
            return
        doc.add_paragraph()
        try:
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_para.add_run().add_picture(BytesIO(png), width=Cm(15))
        except Exception:
            pass
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(f"图: {cap}")
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    while i < len(lines):
        line = lines[i]

        # 一级标题
        if re.match(r'^#\s+', line):
            _para(line[2:].strip(), bold=True, size=18,
                  space_before=12, space_after=6)
            i += 1; continue

        # 二级标题
        if re.match(r'^##\s+', line):
            _para(line[3:].strip(), bold=True, size=14,
                  space_before=8, space_after=4)
            i += 1; continue

        # 三级标题
        if re.match(r'^###\s+', line):
            h3_text = _EMOJI_STRIP_RE.sub('', line[4:].strip()).strip()
            # 若 h3 本身就是"图表N：xxx"标题，消费后续 bullet 并注入图表
            if _WORD_CHART_RE.match(h3_text):
                _para(h3_text, bold=True, size=12, space_before=6, space_after=2)
                i += 1
                while i < len(lines):
                    bl = lines[i].strip()
                    if re.match(r'^[-*]\s+', bl):
                        content = re.sub(r'^[-*]\s+', '', bl)
                        p = doc.add_paragraph(style="List Bullet")
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.paragraph_format.space_after = Pt(2)
                        bparts = re.split(r'(\*\*[^*]+\*\*)', content)
                        for bp in bparts:
                            bm = re.match(r'\*\*([^*]+)\*\*', bp)
                            run = p.add_run(bm.group(1) if bm else bp)
                            run.bold = bool(bm)
                            run.font.size = Pt(10)
                        i += 1
                    elif bl == '':
                        j = i + 1
                        while j < len(lines) and not lines[j].strip():
                            j += 1
                        if j < len(lines) and re.match(r'^[-*]\s+', lines[j].strip()):
                            i += 1
                        else:
                            break
                    else:
                        break
                _word_inject_chart(_WORD_CHART_RE.sub('', h3_text).strip())
                continue
            _para(h3_text, bold=True, size=12, space_before=6)
            i += 1; continue

        # 四级及以下标题 (####) → 去除 emoji，渲染为加粗正文（同 h3 风格）
        if re.match(r'^#{4,}\s+', line):
            h4_text = _EMOJI_STRIP_RE.sub('', re.sub(r'^#{4,}\s+', '', line).strip()).strip()
            _para(h4_text, bold=True, size=11, space_before=4, space_after=2)
            i += 1; continue

        # 水平线
        if re.match(r'^-{3,}\s*$', line):
            _add_hr(doc)
            i += 1; continue

        # Markdown 表格（检测到 | 开头且下一行含 ---）
        if line.startswith("|") and i + 1 < len(lines) and re.match(r'^\|[-| :]+\|', lines[i + 1]):
            # 收集表格行
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith("|"):
                table_lines.append(lines[j])
                j += 1
            # 解析表头 + 数据行（跳过分隔行）
            header_row = [c.strip() for c in table_lines[0].strip('|').split('|')]
            data_rows  = []
            for tl in table_lines[2:]:
                if re.match(r'^\|[-| :]+\|', tl):
                    continue
                data_rows.append([c.strip() for c in tl.strip('|').split('|')])

            n_cols = len(header_row)
            t = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
            t.style = "Table Grid"

            # 表头行
            for ci, cell_text in enumerate(header_row):
                cell = t.rows[0].cells[ci]
                cell.text = cell_text
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
                run.bold = True
                run.font.size = Pt(10)
                # 灰色底色
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9D9D9")
                tc_pr.append(shd)

            # 数据行
            for ri, row_data in enumerate(data_rows, start=1):
                for ci, cell_text in enumerate(row_data[:n_cols]):
                    cell = t.rows[ri].cells[ci]
                    cell.text = cell_text
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.size = Pt(10)

            doc.add_paragraph()  # 表后空行
            i = j; continue

        # 项目符号（- 或 * 开头）；合并软换行续行为单一 bullet 段落
        if re.match(r'^[-*]\s+', line):
            content_parts = [re.sub(r'^[-*]\s+', '', line.strip())]
            i += 1
            while i < len(lines) and not _is_block_start(lines[i]):
                content_parts.append(lines[i].strip())
                i += 1
            content = ''.join(content_parts)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            parts = re.split(r'(\*\*[^*]+\*\*)', content)
            for part in parts:
                m = re.match(r'\*\*([^*]+)\*\*', part)
                run = p.add_run(m.group(1) if m else part)
                run.bold = bool(m)
                run.font.size = Pt(11)
            continue

        # 数字编号列表（1. 2. 等）；合并软换行续行
        if re.match(r'^\d+\.\s+', line):
            content_parts = [re.sub(r'^\d+\.\s+', '', line.strip())]
            i += 1
            while i < len(lines) and not _is_block_start(lines[i]):
                content_parts.append(lines[i].strip())
                i += 1
            content = ''.join(content_parts)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            parts = re.split(r'(\*\*[^*]+\*\*)', content)
            for part in parts:
                m = re.match(r'\*\*([^*]+)\*\*', part)
                run = p.add_run(m.group(1) if m else part)
                run.bold = bool(m)
                run.font.size = Pt(11)
            continue

        # 图表标签行（如 "图表1:xxx"）—— 先输出标签文字，再消费 > 解读块，最后内联插图
        if _WORD_CHART_RE.match(line.strip()):
            _para(line.strip(), size=11, space_before=4, space_after=2)
            i += 1
            # 消费 > blockquote 行：遇空行时探测前方，若仍有 > 则继续
            while i < len(lines):
                bq = lines[i].strip()
                if bq.startswith('>'):
                    _para(bq.lstrip('>').strip(), size=10, space_before=0,
                          space_after=3, color=(75, 85, 99))
                    i += 1
                elif bq == '':
                    j = i + 1
                    while j < len(lines) and lines[j].strip() == '':
                        j += 1
                    if j < len(lines) and lines[j].strip().startswith('>'):
                        i += 1   # 跳过中间空行继续
                    else:
                        break
                else:
                    break
            caption_hint = _WORD_CHART_RE.sub('', line.strip()).strip()
            _word_inject_chart(caption_hint)
            continue

        # Markdown 图片占位符 ![图表N:caption]() → 注入匹配图表（不输出文字）
        if _MD_IMG_RE.match(line.strip()):
            raw_cap = _MD_IMG_RE.match(line.strip()).group(1)
            caption_hint = re.sub(r'^图表?\s*\d+\s*[：:]', '', raw_cap).strip()
            _word_inject_chart(caption_hint)
            i += 1; continue

        # blockquote（> 行，非图表标签后的独立解读）
        if line.strip().startswith('>'):
            _para(line.strip().lstrip('>').strip(), size=10, space_before=0,
                  space_after=3, color=(75, 85, 99))
            i += 1; continue

        # 空行
        if not line.strip():
            i += 1; continue

        # 普通正文（支持内联加粗）；合并连续软换行行为单一段落
        para_parts = [line.strip()]
        i += 1
        while i < len(lines) and not _is_block_start(lines[i]):
            para_parts.append(lines[i].strip())
            i += 1
        _para(''.join(para_parts), size=11, space_before=0, space_after=4)

    # ── 未匹配图表兜底：追加到末尾 ───────────────────────────────────────────
    if _word_unplaced and any(png for _, png in _word_unplaced):
        doc.add_paragraph()
        _add_hr(doc)
        for cap, png in _word_unplaced:
            if png is None:
                continue
            doc.add_paragraph()
            try:
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_para.add_run().add_picture(BytesIO(png), width=Cm(15))
            except Exception:
                pass
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(f"图: {cap}")
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 函数 2: export_chart_as_png ─────────────────────────────────────────────

_SYSTEM_CHROME_PATHS = [
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
]

def _find_system_chrome() -> str | None:
    """返回系统 Chrome 可执行文件路径，未找到返回 None。"""
    import os
    for p in _SYSTEM_CHROME_PATHS:
        if os.path.isfile(p):
            return p
    return None


def _kaleido_major() -> int:
    """返回已安装 kaleido 的主版本号（0 或 1）。"""
    try:
        import importlib.metadata
        ver = importlib.metadata.version("kaleido")
        return int(ver.split(".")[0])
    except Exception:
        return 0


def export_chart_as_png(fig, width: int = 960, height: int = 440,
                        timeout: int = 30) -> bytes | None:
    """
    把 Plotly Figure 导出为 PNG bytes，自动适配 kaleido v0 / v1。

    - kaleido v1: 用 calc_fig_sync + 系统 Chrome 路径（跳过 choreographer 自动下载）
    - kaleido v0: 用 plotly.io.to_image（内置可执行文件，不需要 Chrome）
    用 daemon 线程 + threading.Event 保护，超时返回 None，永不阻塞调用方。
    """
    import threading

    ver = _kaleido_major()
    chrome_path = _find_system_chrome()
    print(f"[export_chart_as_png] kaleido v{ver}, Chrome={chrome_path or 'N/A'}")

    result = [None]
    done   = threading.Event()

    def _worker():
        try:
            if ver >= 1:
                import kaleido as _kaleido
                fig_dict = fig.to_dict()
                kopts = {"path": chrome_path} if chrome_path else {}
                result[0] = _kaleido.calc_fig_sync(
                    fig_dict,
                    opts={"format": "png", "width": width, "height": height, "scale": 1.5},
                    kopts=kopts,
                )
            else:
                # kaleido v0: use plotly's built-in to_image
                import plotly.io as pio
                result[0] = pio.to_image(fig, format="png", width=width, height=height, scale=1.5)
            print(f"[export_chart_as_png] 成功 size={len(result[0])} bytes")
        except Exception as e:
            print(f"[export_chart_as_png] 异常: {e}")
        finally:
            done.set()

    t = threading.Thread(target=_worker, daemon=True)
    t.start()
    if done.wait(timeout=timeout):
        return result[0]
    print(f"[export_chart_as_png] 超时（>{timeout}s），跳过此图表")
    return None


# ── 函数 3: build_html_report ────────────────────────────────────────────────

# 专业报告 CSS（深海蓝 #1B3A6B + 橙色 #E87722，微软雅黑优先）
_HTML_REPORT_CSS = """
@page {
    size: A4;
    margin: 20mm 22mm 25mm 22mm;
    @top-center {
        content: element(header);
        vertical-align: bottom;
    }
    @bottom-left {
        content: "Confidential";
        font-size: 8pt;
        color: #9CA3AF;
        font-family: 'Microsoft YaHei', sans-serif;
    }
    @bottom-right {
        content: "第 " counter(page) " 页 / 共 " counter(pages) " 页";
        font-size: 8pt;
        color: #6B7280;
        font-family: 'Microsoft YaHei', sans-serif;
    }
}

#page-header {
    position: running(header);
    width: 100%;
    border-bottom: 2px solid #1B3A6B;
    padding-bottom: 4pt;
    display: flex;
    justify-content: space-between;
    align-items: center;
}
#page-header .header-title {
    font-size: 8.5pt;
    color: #1B3A6B;
    font-weight: 600;
    font-family: 'Microsoft YaHei', sans-serif;
}
#page-header .header-date {
    font-size: 8pt;
    color: #9CA3AF;
    font-family: 'Microsoft YaHei', sans-serif;
}

* { box-sizing: border-box; margin: 0; padding: 0; }

body {
    font-family: 'Microsoft YaHei', 'PingFang SC', 'STHeiti',
                 'WenQuanYi Micro Hei', 'Noto Sans CJK SC', sans-serif;
    font-size: 10.5pt;
    line-height: 1.75;
    color: #2C2C2C;
    background: white;
}

.cover {
    text-align: center;
    padding: 60pt 30pt 40pt 30pt;
    border-bottom: 4px solid #1B3A6B;
    margin-bottom: 36pt;
    page-break-after: always;
}
.cover-tag {
    display: inline-block;
    background: #E87722;
    color: white;
    font-size: 8pt;
    font-weight: 700;
    letter-spacing: 2px;
    padding: 3pt 10pt;
    border-radius: 2pt;
    margin-bottom: 20pt;
    text-transform: uppercase;
}
.cover-title {
    font-size: 24pt;
    font-weight: 700;
    color: #1B3A6B;
    line-height: 1.3;
    margin-bottom: 16pt;
}
.cover-subtitle {
    font-size: 11pt;
    color: #6B7280;
    margin-bottom: 32pt;
}
.cover-meta {
    font-size: 9pt;
    color: #9CA3AF;
    border-top: 1px solid #E5E7EB;
    padding-top: 16pt;
}
.cover-meta span { margin: 0 12pt; }

h1 {
    font-size: 16pt;
    font-weight: 700;
    color: #1B3A6B;
    margin-top: 30pt;
    margin-bottom: 12pt;
    padding-bottom: 6pt;
    border-bottom: 2.5px solid #1B3A6B;
    page-break-after: avoid;
}
h1 .section-num {
    display: inline-block;
    background: #1B3A6B;
    color: white;
    font-size: 9pt;
    padding: 1pt 7pt;
    border-radius: 2pt;
    margin-right: 8pt;
    vertical-align: middle;
}
h2 {
    font-size: 13pt;
    font-weight: 700;
    color: #1B3A6B;
    margin-top: 22pt;
    margin-bottom: 8pt;
    padding-left: 10pt;
    border-left: 4px solid #E87722;
    page-break-after: avoid;
}
h3 {
    font-size: 11pt;
    font-weight: 600;
    color: #374151;
    margin-top: 16pt;
    margin-bottom: 6pt;
    page-break-after: avoid;
}
p {
    margin-bottom: 9pt;
    text-align: justify;
    orphans: 3;
    widows: 3;
}
strong { color: #1B3A6B; font-weight: 700; }
ul, ol { margin: 8pt 0 10pt 18pt; }
li { margin-bottom: 4pt; line-height: 1.7; }
li::marker { color: #E87722; }

table {
    width: 100%;
    border-collapse: collapse;
    margin: 14pt 0;
    font-size: 9.5pt;
    page-break-inside: avoid;
}
thead tr { background: #1B3A6B; color: white; }
thead th {
    padding: 8pt 10pt;
    text-align: left;
    font-weight: 600;
    font-size: 9pt;
    letter-spacing: 0.3px;
}
thead th:first-child, tbody td:first-child { text-align: center; }
tbody tr:nth-child(even) { background: #F5F7FA; }
tbody tr:hover { background: #EFF6FF; }
tbody td {
    padding: 7pt 10pt;
    border-bottom: 1px solid #E5E7EB;
    color: #374151;
}
tbody td:not(:first-child):not(:nth-child(2)) { text-align: right; }
tbody tr.highlight td { font-weight: 700; color: #1B3A6B; }
.table-caption {
    font-size: 9pt; color: #6B7280;
    text-align: center; margin-top: -8pt; margin-bottom: 12pt; font-style: italic;
}

.chart-container {
    margin: 16pt 0;
    text-align: center;
    page-break-inside: avoid;
}
.chart-container img {
    max-width: 100%;
    border: 1px solid #E5E7EB;
    border-radius: 4pt;
    box-shadow: 0 1px 4px rgba(0,0,0,0.08);
}
.chart-caption {
    font-size: 8.5pt; color: #6B7280;
    margin-top: 6pt; font-style: italic; text-align: center;
}
.chart-number { font-weight: 700; color: #1B3A6B; }

.page-break { page-break-before: always; }
.no-break { page-break-inside: avoid; }

.appendix {
    border-top: 2px solid #E5E7EB;
    margin-top: 30pt;
    padding-top: 20pt;
}
.appendix h1 { color: #6B7280; border-bottom-color: #E5E7EB; font-size: 14pt; }

.watermark {
    position: fixed;
    top: 50%; left: 50%;
    transform: translate(-50%, -50%) rotate(-35deg);
    font-size: 60pt;
    color: rgba(27, 58, 107, 0.04);
    font-weight: 900;
    pointer-events: none;
    z-index: -1;
    white-space: nowrap;
}
"""


def _chart_html(chart_n: int, b64: str, caption: str) -> str:
    """生成单张图表的 HTML 片段。"""
    return (
        f'<div class="chart-container">'
        f'<img src="data:image/png;base64,{b64}" alt="{caption}"/>'
        f'<p class="chart-caption">'
        f'<span class="chart-number">图{chart_n}</span> {caption}'
        f'</p></div>'
    )


def _match_chart_for_section(section_title: str, unplaced: list) -> dict | None:
    """
    三级优先级为章节标题匹配图表。
    P1: 图表标题直接包含章节关键词（bigram / 英文词）。
    P2: 章节与图表标题 bigram 集合有交集。
    P3: 分析类单字兜底（析/层/分/趋/比/额/率/布）——解决"三、详细分析" vs "用户分层" 等典型不匹配。
    """
    # 分析类关键单字：出现在这些字的章节可以接纳任何分析图表
    _ANALYSIS_CH = frozenset('析层分趋比额率布统量增跌变')

    def keywords(s: str) -> set:
        result = set()
        for w in re.findall(r'[a-zA-Z]{3,}', s):
            result.add(w.lower())
        for run in re.findall(r'[一-鿿]+', s):
            for i in range(len(run) - 1):
                result.add(run[i:i + 2])
        return result

    sec_kw = keywords(section_title)
    sec_ch = set(re.findall(r'[一-鿿]', section_title))

    # P1: 直接字符串包含
    for c in unplaced:
        if not c.get('placed'):
            for w in sec_kw:
                if len(w) >= 2 and w in c['caption']:
                    c['placed'] = True
                    return c

    # P2: bigram 集合交集
    for c in unplaced:
        if not c.get('placed'):
            if sec_kw & keywords(c['caption']):
                c['placed'] = True
                return c

    # P3: 分析类单字兜底（章节含分析类字 AND 图表也含分析类字）
    sec_analysis = sec_ch & _ANALYSIS_CH
    if sec_analysis:
        for c in unplaced:
            if not c.get('placed'):
                cap_ch = set(re.findall(r'[一-鿿]', c['caption']))
                if sec_analysis & cap_ch & _ANALYSIS_CH:
                    c['placed'] = True
                    return c

    return None


def build_html_report(
    markdown_content: str,
    title: str,
    chart_figures: list = None,
    chart_captions: list = None,
    generated_date: str = None,
) -> str:
    """
    把 Markdown 报告内容 + Plotly 图表合并为专业排版的 HTML。

    图表匹配规则：
    - 按章节（## 标题）关键词匹配，优先完全包含，其次关键词交集
    - 匹配到的图表嵌入对应章节末尾；未匹配的图表汇总到附图附录
    """
    import markdown as md_lib

    chart_figures  = chart_figures  or []
    chart_captions = chart_captions or []
    if not generated_date:
        generated_date = datetime.now().strftime("%Y年%m月%d日 %H:%M")

    # ── Step 1: 导出所有图表为 base64 PNG ──────────────────────────────────────
    chart_data = []
    for i, (fig, cap) in enumerate(zip(chart_figures, chart_captions)):
        png_bytes = export_chart_as_png(fig, width=1100, height=500)
        if png_bytes:
            chart_data.append({
                'index':   i + 1,
                'caption': cap,
                'b64':     base64.b64encode(png_bytes).decode(),
                'placed':  False,
            })

    # ── Step 2: Markdown → HTML ────────────────────────────────────────────────
    body_html = md_lib.markdown(
        markdown_content,
        extensions=['tables', 'fenced_code', 'nl2br'],
    )

    # ── Step 3: 按 <h2> 章节分段，匹配并插入图表 ─────────────────────────────
    # 用 re.split 把 body_html 按 <h2>...</h2> 切割
    # 结果格式：[before_h2, h2_tag, section_content, h2_tag, section_content, ...]
    parts = re.split(r'(<h2>[^<]*</h2>)', body_html)

    result_parts = []
    chart_n = 0  # 全局图表编号
    unplaced = [c for c in chart_data]  # 尚未放置的图表（共享引用）

    for idx, part in enumerate(parts):
        result_parts.append(part)
        m = re.match(r'<h2>([^<]*)</h2>', part)
        if not m:
            continue
        section_title = m.group(1)
        # 在 unplaced 里找本章节图表，插入到下一个 h2 之前（即当前 section_content 末尾）
        # section_content 在 parts[idx+1]（如果存在）
        # 不修改 parts[idx+1]；改为在 parts[idx+1] 之后插入图表 HTML
        matched = _match_chart_for_section(section_title, [c for c in unplaced if not c['placed']])
        if matched:
            chart_n += 1
            matched['placed'] = True
            matched['chart_n'] = chart_n  # 更新实际编号
            # 将图表注入紧跟在该章节内容之后
            # 找到本 h2 对应的 section_content（parts[idx+1]），在其末尾插入
            if idx + 1 < len(parts):
                result_parts.append(parts[idx + 1])   # section content
                result_parts.append(_chart_html(chart_n, matched['b64'], matched['caption']))
                # 跳过 parts[idx+1]，否则会被再次 append 在下一轮
                parts[idx + 1] = ''  # 已处理，清空防止重复
            else:
                result_parts.append(_chart_html(chart_n, matched['b64'], matched['caption']))

    body_with_charts = ''.join(result_parts)

    # ── Step 4: 未匹配图表放附录 ───────────────────────────────────────────────
    still_unplaced = [c for c in chart_data if not c['placed']]
    appendix_html = ''
    if still_unplaced:
        appendix_html = '<div class="appendix"><h1>附图</h1>'
        for c in still_unplaced:
            chart_n += 1
            appendix_html += _chart_html(chart_n, c['b64'], c['caption'])
        appendix_html += '</div>'

    # ── Step 5: 组装完整 HTML ──────────────────────────────────────────────────
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>{_HTML_REPORT_CSS}</style>
</head>
<body>
    <div class="watermark">CONFIDENTIAL</div>
    <div id="page-header">
        <span class="header-title">{title}</span>
        <span class="header-date">{generated_date}</span>
    </div>
    <div class="cover">
        <div class="cover-tag">Analysis Report</div>
        <div class="cover-title">{title}</div>
        <div class="cover-subtitle">AI 数据分析 Agent 自动生成</div>
        <div class="cover-meta">
            <span>生成时间: {generated_date}</span>
            <span>数据来源: Olist Brazilian E-Commerce</span>
        </div>
    </div>
    {body_with_charts}
    {appendix_html}
</body>
</html>"""


# ── 函数 4a: _build_pdf_safe_html（PDF 专用简化 HTML）────────────────────────

def _build_pdf_safe_html(markdown_content: str, title: str) -> str:
    """
    生成 xhtml2pdf 兼容的 HTML。
    与 build_html_report 的区别：
    - 无 @page 命名 margin boxes（@top-right 等）—— xhtml2pdf 不支持，会卡死
    - 无 :nth-child 伪选择器 —— xhtml2pdf 不支持
    - 不用 @font-face file:// URI —— 由调用方通过 reportlab API 预注册字体
    """
    import markdown as md_lib

    body_html = md_lib.markdown(
        markdown_content,
        extensions=["tables", "fenced_code"],
    )
    generated_date = datetime.now().strftime("%Y-%m-%d %H:%M")

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8"/>
<title>{title}</title>
<style>
@page {{ size: A4; margin: 20mm; }}
body   {{ font-family: Arial, sans-serif; font-size: 11pt;
          line-height: 1.6; color: #333; }}
h1     {{ font-size: 18pt; font-weight: bold; color: #1a1a2e;
          border-bottom: 2px solid #4C72B0; padding-bottom: 4pt;
          margin-top: 0; text-align: center; }}
h2     {{ font-size: 14pt; font-weight: bold; color: #1a1a2e; margin-top: 14pt; }}
h3     {{ font-size: 12pt; font-weight: bold; margin-top: 10pt; }}
p      {{ margin: 6pt 0; }}
table  {{ width: 100%; border-collapse: collapse; margin: 10pt 0; font-size: 10pt; }}
th     {{ background-color: #F2F2F2; font-weight: bold; text-align: left;
          border: 1px solid #CCC; padding: 5px 8px; }}
td     {{ border: 1px solid #CCC; padding: 5px 8px; }}
ul, ol {{ padding-left: 1.5em; margin: 6pt 0; }}
li     {{ margin: 3pt 0; }}
code   {{ background: #F5F5F5; padding: 2px 4px; font-size: 9pt; }}
pre    {{ background: #F5F5F5; padding: 8px; font-size: 9pt; }}
blockquote {{ border-left: 4px solid #4C72B0; margin: 8pt 0;
              padding: 4pt 10pt; color: #555; }}
hr     {{ border: none; border-top: 1px solid #CCC; margin: 12pt 0; }}
.rpt-date {{ text-align: center; color: #888; font-size: 9pt; margin-bottom: 8pt; }}
</style>
</head>
<body>
<h1>{title}</h1>
<p class="rpt-date">生成时间：{generated_date}</p>
<hr/>
{body_html}
</body>
</html>"""


# ── 函数 4b & 5: markdown_to_pdf（reportlab 直接生成）────────────────────────

# 候选中文字体路径（按优先级）；.ttc 用 subfontIndex=0 加载
_CJK_FONT_PATHS = [
    r"C:\Windows\Fonts\msyh.ttc",       # 微软雅黑 Regular（Win10/11，字形最全）
    r"C:\Windows\Fonts\msyh.ttf",       # 部分系统为 TTF 格式
    r"C:\Windows\Fonts\msyhbd.ttc",     # 微软雅黑 Bold（备用）
    r"C:\Windows\Fonts\simhei.ttf",     # 黑体（旧系统备选）
    r"C:\Windows\Fonts\simkai.ttf",
    r"C:\Windows\Fonts\simfang.ttf",
    r"C:\Windows\Fonts\STZHONGS.TTF",
]


def _register_cjk_font() -> tuple:
    """
    注册第一个可用的 CJK 字体（TTF/TTC），返回 (字体名, 字体路径)。
    找不到可用字体时返回 ('Helvetica', '')。
    """
    import os
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for path in _CJK_FONT_PATHS:
        if not os.path.exists(path):
            continue
        font_name = "CJKFont"
        if font_name not in pdfmetrics.getRegisteredFontNames():
            try:
                if path.lower().endswith(".ttc"):
                    pdfmetrics.registerFont(TTFont(font_name, path, subfontIndex=0))
                else:
                    pdfmetrics.registerFont(TTFont(font_name, path))
            except Exception:
                continue
        return font_name, path
    return "Helvetica", ""


def markdown_to_pdf(
    markdown_content: str,
    title: str,
    chart_figures: list = None,
    chart_captions: list = None,
    chart_png_bytes: list = None,
) -> bytes:
    """
    Markdown → PDF，使用 reportlab 直接生成（不经过 xhtml2pdf/HTML）。
    中文字体通过 pdfmetrics.registerFont 嵌入，彻底解决乱码问题。
    chart_png_bytes: 可选，预先导出的 PNG bytes 列表（与 chart_figures 等长）。
    传入时直接使用，跳过 export_chart_as_png，避免在 60s 超时内串行导出超时。
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor, white
    pt = 1
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak, KeepTogether,
    )

    FONT, _font_path = _register_cjk_font()
    _is_yahe = "msyh" in _font_path.lower()

    # ── 颜色（与 HTML CSS 完全对齐）──────────────────────────────────────────
    C_NAVY   = HexColor("#1B3A6B")   # 深海蓝：标题、边框、表头
    C_ORANGE = HexColor("#E87722")   # 橙色：h2 左边框、badge
    C_DARK   = HexColor("#2C2C2C")   # 正文深灰
    C_GRAY   = HexColor("#6B7280")   # 辅助中灰
    C_LIGHT  = HexColor("#F5F7FA")   # 偶数行背景
    C_BDR    = HexColor("#E5E7EB")   # 边框浅灰
    C_META   = HexColor("#9CA3AF")   # 封面元数据灰

    def S(name, **kw):
        return ParagraphStyle(name, fontName=FONT, **kw)

    # ── 样式定义（对照 HTML CSS）─────────────────────────────────────────────
    # 封面
    s_badge    = S("BDG", fontSize=8,  textColor=white,   alignment=1,
                   leading=10,  spaceAfter=0)
    s_title    = S("T",   fontSize=24, textColor=C_NAVY,  alignment=1,
                   leading=30,  spaceBefore=0, spaceAfter=10*pt)
    s_subtitle = S("SUB", fontSize=11, textColor=C_GRAY,  alignment=1,
                   leading=16,  spaceAfter=24*pt)
    s_meta     = S("MTA", fontSize=9,  textColor=C_META,  alignment=1,
                   leading=13,  spaceAfter=0)
    # 正文
    s_date     = S("D",   fontSize=9,  textColor=C_GRAY,  alignment=1,
                   leading=13,  spaceAfter=6*pt, wordWrap='CJK')
    s_h1       = S("H1",  fontSize=16, textColor=C_NAVY,
                   spaceBefore=22*pt, spaceAfter=6*pt, leading=22, wordWrap='CJK')
    s_h2_text  = S("H2T", fontSize=13, textColor=C_NAVY,
                   spaceBefore=0, spaceAfter=0, leading=20, wordWrap='CJK')
    s_h3       = S("H3",  fontSize=11, textColor=HexColor("#374151"),
                   spaceBefore=9*pt, spaceAfter=3*pt, leading=16, wordWrap='CJK')
    s_body     = S("B",   fontSize=10, textColor=C_DARK,
                   spaceAfter=6*pt, leading=17, wordWrap='CJK')
    s_bullet   = S("BL",  fontSize=10, textColor=C_DARK,
                   spaceAfter=3*pt, leading=16, leftIndent=14*pt, wordWrap='CJK')
    s_num      = S("NL",  fontSize=10, textColor=C_DARK,
                   spaceAfter=3*pt, leading=16, leftIndent=18*pt, firstLineIndent=-18*pt, wordWrap='CJK')
    s_cell     = S("C",   fontSize=9,  textColor=HexColor("#374151"), leading=14)
    s_hcell    = S("HC",  fontSize=9,  textColor=white,  leading=14)

    def _h2_flowable(text: str) -> Table:
        """H2：用 Table + LINEBEFORE 模拟 CSS border-left: 4px solid #E87722。"""
        cell = Paragraph(text, s_h2_text)
        tbl  = Table([[cell]], colWidths=[avail_w])
        tbl.setStyle(TableStyle([
            ("LEFTPADDING",    (0, 0), (-1, -1), 12),
            ("RIGHTPADDING",   (0, 0), (-1, -1), 0),
            ("TOPPADDING",     (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 5),
            ("LINEBEFORE",     (0, 0), (0,  -1), 4, C_ORANGE),
        ]))
        return tbl

    # ── 符号安全转换 ─────────────────────────────────────────────────────────
    # 层1：不管用哪种字体都要转（XML 安全 + 全角半角统一）
    _SYM_ALWAYS = {
        '￥': '¥',      # ￿e5(full-width) -> ¥ (half-width yen/yuan)
        '‘': "'",         # ‘ left single curly quote
        '’': "'",         # ’ right single curly quote
        '“': '"',         # “ left double curly quote
        '”': '"',         # ” right double curly quote
        '–': '-',          # en dash
        '—': '--',         # em dash
        '…': '...',        # ellipsis
        **{chr(0xFF01 + i): chr(0x21 + i) for i in range(94)},  # full-width ASCII -> half
    }
    _SYM_FALLBACK = {   # only applied when font is NOT YaHei (SimHei fallback)
        '×':   'x',     # x (multiplication)
        '÷':   '/',     # / (division)
        '±':   '+/-',   # +/- (plus-minus)
        '≥': '>=',    # >=
        '≤': '<=',    # <=
        '≠': '!=',    # !=
        '≈': '~=',    # approximately
        '√': 'sqrt',  # sqrt
        '∞': 'inf',   # infinity
        '∑': 'sum',   # sum
        '→': '->',    # right arrow
        '←': '<-',    # left arrow
        '↑': '^',     # up arrow
        '↓': 'v',     # down arrow
        '▲': '(+)',   # up triangle
        '▼': '(-)',   # down triangle
        '©':   '(c)',   # copyright
        '®':   '(R)',   # registered
        '™': '(TM)',  # trademark
        '·':   '.',     # middle dot
        **{chr(0x2460 + i): str(i + 1) for i in range(20)},  # circled 1-20
    }
    _SYM_MAP = dict(_SYM_ALWAYS)
    if not _is_yahe:
        _SYM_MAP.update(_SYM_FALLBACK)

    # ── 千位符：匹配数字（可带小数）后紧跟"元"，或整数部分 ≥ 4 位的金额数字 ──
    _MONEY_RE = re.compile(
        r'(\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d+(?:\.\d+)?)'   # 数字（含已有逗号）
        r'(?=\s*元)'                                           # 后接"元"
    )
    _LARGE_INT_RE = re.compile(r'\b(\d{4,})(\.\d+)?\b')     # ≥4位整数（不跟年份检测）

    def _fmt_num(m: re.Match) -> str:
        """给数字加千位符。"""
        raw = m.group(1).replace(',', '')
        dec  = m.group(2) or ''
        try:
            n = int(raw)
            return f"{n:,}{dec}"
        except ValueError:
            return m.group(0)

    def _add_thousands(text: str) -> str:
        """对"数字+元"和独立大整数（≥1000，含小数）添加千位符。"""
        # 先处理"数字元"
        def _repl_yuan(m):
            raw = m.group(1).replace(',', '')
            try:
                if '.' in raw:
                    ip, dp = raw.split('.', 1)
                    return f"{int(ip):,}.{dp}"
                return f"{int(raw):,}"
            except ValueError:
                return m.group(0)
        text = _MONEY_RE.sub(_repl_yuan, text)
        # 再处理独立大数（≥4位整数，且不像年份：不匹配 1900-2099 四位纯整数）
        def _repl_large(m):
            raw, dec = m.group(1), m.group(2) or ''
            if len(raw) == 4 and 1900 <= int(raw) <= 2099:
                return m.group(0)  # 年份不加千位符
            try:
                n = int(raw)
                if n < 1000:
                    return m.group(0)
                return f"{n:,}{dec}"
            except ValueError:
                return m.group(0)
        text = _LARGE_INT_RE.sub(_repl_large, text)
        return text

    def inline(text: str) -> str:
        """千位符 + XML 转义 + 符号规范化 + **bold** → <b>。"""
        text = _EMOJI_STRIP_RE.sub('', text).strip()   # CJK 字体无法渲染的 emoji 直接去除
        text = _add_thousands(text)
        for src, dst in _SYM_MAP.items():
            text = text.replace(src, dst)
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # 数字与中文量词/单位之间插入不换行空格，防止"3.2\n倍"等不合理断行
        text = re.sub(r'(\d+(?:\.\d+)?)\s*([倍个万亿千百元笔次年月日%])', r'\1&nbsp;\2', text)
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        return text

    buf  = BytesIO()
    doc  = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=20*mm, leftMargin=20*mm,
        topMargin=22*mm,   bottomMargin=20*mm,
        title=title,
    )
    story = []
    generated_date = datetime.now().strftime("%Y-%m-%d %H:%M")

    lines   = markdown_content.split("\n")
    i       = 0
    avail_w = A4[0] - 40*mm

    # ── 封面页（对照 HTML cover 设计）────────────────────────────────────────
    # "ANALYSIS REPORT" 橙色 badge（居中 Table）
    badge_tbl = Table([[Paragraph("ANALYSIS REPORT", s_badge)]],
                      colWidths=[70*mm])
    badge_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), C_ORANGE),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
    ]))
    badge_tbl.hAlign = "CENTER"

    cover = KeepTogether([
        Spacer(1, 50*pt),
        badge_tbl,
        Spacer(1, 18*pt),
        Paragraph(inline(title), s_title),
        Paragraph("AI 数据分析 Agent 自动生成", s_subtitle),
        HRFlowable(width="100%", thickness=1, color=C_BDR, spaceAfter=12*pt),
        Paragraph(
            f"生成时间：{generated_date}　　数据来源：Olist Brazilian E-Commerce",
            s_meta,
        ),
        Spacer(1, 36*pt),
        HRFlowable(width="100%", thickness=4, color=C_NAVY),
    ])
    story.append(cover)
    story.append(PageBreak())

    # ── 图表嵌入：预计算 section→chart 映射 ───────────────────────────────────
    _chart_figs  = chart_figures   or []
    _chart_caps  = chart_captions  or []
    _chart_bytes = chart_png_bytes or [None] * len(_chart_figs)
    _ok_pngs = sum(1 for b in _chart_bytes if b is not None)
    print(f"[markdown_to_pdf] 收到图表: {len(_chart_figs)} 张, PNG 有效: {_ok_pngs}/{len(_chart_bytes)}")
    _pdf_charts  = [
        {'caption': cap, 'fig': fig, 'png': png, 'placed': False}
        for fig, cap, png in zip(_chart_figs, _chart_caps, _chart_bytes)
    ]
    _sec_titles = re.findall(r'^## +(.+)$', markdown_content, re.MULTILINE)
    print(f"[markdown_to_pdf] 报告章节(##): {_sec_titles}")
    # _sec_chart_map: section_idx → list[chart_dict]，一个章节可嵌入多张图
    _sec_chart_map: dict = {}
    for _si, _st in enumerate(_sec_titles):
        # 贪婪匹配：一直取，直到没有新匹配（同一章节可放多张相关图）
        while True:
            _candidates = [c for c in _pdf_charts if not c['placed']]
            if not _candidates:
                break
            _mc = _match_chart_for_section(_st, _candidates)
            if not _mc:
                break
            _mc['placed'] = True  # 防止同一图表被重复匹配到多个章节
            _sec_chart_map.setdefault(_si, []).append(_mc)
    for _si, _st in enumerate(_sec_titles):
        _matched = [c['caption'] for c in _sec_chart_map.get(_si, [])]
        print(f"[markdown_to_pdf] 章节[{_si}]'{_st}' → 匹配图表: {_matched or '无'}")
    print(f"[markdown_to_pdf] 未匹配图表（进附录）: {[c['caption'] for c in _pdf_charts if not c['placed']] or '无'}")
    # 预计算只做分组，inline 注入时才真正设置 placed；重置标记供正式排版使用
    for _c in _pdf_charts:
        _c['placed'] = False
    _cur_sec = [-1]

    # ── ### 子节级图表预匹配 ─────────────────────────────────────────────────
    _h3_titles = re.findall(r'^###\s+(.+)$', markdown_content, re.MULTILINE)
    print(f"[markdown_to_pdf] 报告子节(###): {_h3_titles}")
    _h3_chart_map: dict = {}
    for _hi, _ht in enumerate(_h3_titles):
        _candidates = [c for c in _pdf_charts if not c['placed']]
        if not _candidates:
            continue
        _mc = _match_chart_for_section(_ht, _candidates)
        if _mc:
            _mc['placed'] = True
            _h3_chart_map[_hi] = [_mc]
    for _hi, _ht in enumerate(_h3_titles):
        _matched = [c['caption'] for c in _h3_chart_map.get(_hi, [])]
        print(f"[markdown_to_pdf] 子节[{_hi}]'{_ht}' → 匹配图表: {_matched or '无'}")
    print(f"[markdown_to_pdf] h3未匹配图表: {[c['caption'] for c in _pdf_charts if not c['placed']] or '无'}")
    for _c in _pdf_charts:
        _c['placed'] = False    # 重置，供正式排版使用
    _cur_h3 = [-1]

    # 图表解读行正则：匹配 "图表1:" / "图1：" / "图表 2:" 等写法
    _CHART_LABEL_RE = re.compile(r'^图表?\s*\d+\s*[：:]')

    def _inject_chart_inline(caption_hint: str):
        """匹配并立即注入一张图表（inline 模式，图表解读后插入）。"""
        from reportlab.platypus import Image as _RLImage
        candidates = [c for c in _pdf_charts if not c.get('placed')]
        if not candidates:
            return
        matched = _match_chart_for_section(caption_hint, candidates)
        if matched is None:
            matched = candidates[0]          # 兜底：取第一张未放置的
        matched['placed'] = True
        png = matched.get('png')
        if png:
            story.append(Spacer(1, 6 * pt))
            story.append(_RLImage(BytesIO(png), width=avail_w, height=avail_w * 440 / 960))
            story.append(Paragraph(f"图: {inline(matched['caption'])}", s_date))
            story.append(Spacer(1, 8 * pt))

    def _add_sec_chart():
        """把当前章节仍未放置的匹配图表兜底插入 story（fallback）。"""
        idx = _cur_sec[0]
        if idx < 0:
            return
        charts_for_sec = _sec_chart_map.get(idx)
        if not charts_for_sec:
            return
        from reportlab.platypus import Image as _RLImage
        for c in charts_for_sec:
            if c.get('placed'):              # inline 已放置则跳过
                continue
            png = c.get('png')
            if png is None:
                continue
            c['placed'] = True
            story.append(Spacer(1, 6 * pt))
            story.append(_RLImage(BytesIO(png), width=avail_w, height=avail_w * 440 / 960))
            story.append(Paragraph(f"图: {inline(c['caption'])}", s_date))
            story.append(Spacer(1, 8 * pt))

    def _add_h3_chart():
        """把当前 ### 子节匹配的图表插入 story（子节末尾注入）。"""
        from reportlab.platypus import Image as _RLImage
        idx = _cur_h3[0]
        if idx < 0:
            return
        for c in _h3_chart_map.get(idx, []):
            if c.get('placed'):
                continue
            png = c.get('png')
            if png is None:
                continue
            c['placed'] = True
            story.append(Spacer(1, 6 * pt))
            story.append(_RLImage(BytesIO(png), width=avail_w,
                                  height=avail_w * 440 / 960))
            story.append(Paragraph(f"图: {inline(c['caption'])}", s_date))
            story.append(Spacer(1, 8 * pt))

    def _handle_chart_label_block(label_text: str, label_style, start_i: int) -> int:
        """
        把图表标签行 + 后续解读文字 + 图表图片打包为 KeepTogether 块。
        消费顺序：> blockquote 或普通段落（直到遇到新块为止），然后紧跟图表图片。
        返回处理后的新行下标。
        """
        from reportlab.platypus import Image as _RLImage
        s_bq = S("BQ", fontSize=10, textColor=HexColor("#4B5563"),
                 leftIndent=14*pt, spaceAfter=4*pt, leading=17, wordWrap='CJK')
        block = [Paragraph(inline(label_text), label_style)]
        ci = start_i

        while ci < len(lines):
            bq = lines[ci].strip()
            if bq.startswith('>'):
                block.append(Paragraph(inline(bq.lstrip('>').strip()), s_bq))
                ci += 1
            elif bq == '':
                j = ci + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and not _is_block_start(lines[j]):
                    ci += 1   # 跳过空行，继续消费解读文字
                else:
                    break
            elif not _is_block_start(bq):   # 普通解读段落
                block.append(Paragraph(inline(bq), s_bq))
                ci += 1
            else:
                break

        # 图表在解读文字之后注入
        caption_hint = _CHART_LABEL_RE.sub('', label_text).strip()
        candidates = [c for c in _pdf_charts if not c.get('placed')]
        if candidates:
            matched = _match_chart_for_section(caption_hint, candidates)
            if matched is None:
                matched = candidates[0]
            matched['placed'] = True
            png = matched.get('png')
            if png:
                block.append(Spacer(1, 6 * pt))
                block.append(_RLImage(BytesIO(png), width=avail_w,
                                      height=avail_w * 440 / 960))
                block.append(Paragraph(f"图: {inline(matched['caption'])}", s_date))
                block.append(Spacer(1, 8 * pt))

        story.append(KeepTogether(block))
        return ci

    while i < len(lines):
        line     = lines[i]
        stripped = line.strip()

        if re.match(r'^# [^#]', stripped):
            story.append(Paragraph(inline(stripped[2:]), s_h1))
            story.append(HRFlowable(width="100%", thickness=2, color=C_NAVY, spaceAfter=5*pt))
            i += 1

        elif re.match(r'^## [^#]', stripped):
            _add_h3_chart()            # 注入当前 h3 子节末尾图表
            _add_sec_chart()           # 把上一章节匹配的图表插在章节末尾（兜底）
            _cur_sec[0] += 1
            _cur_h3[0] = -1            # 新章节时重置 h3 计数器
            story.append(Spacer(1, 14*pt))
            story.append(_h2_flowable(inline(stripped[3:])))
            story.append(Spacer(1, 5*pt))
            i += 1

        elif re.match(r'^### ', stripped):
            _add_h3_chart()            # 注入上一 h3 子节末尾图表
            _cur_h3[0] += 1            # 推进子节计数器
            h3_text = stripped[4:]
            # 若 h3 标题本身是"图表N：xxx"，消费后续 bullet 并立即注入图表
            if _CHART_LABEL_RE.match(h3_text):
                story.append(Paragraph(inline(h3_text), s_h3))
                i += 1
                s_bq2 = S("BQ2", fontSize=10, textColor=HexColor("#4B5563"),
                          leftIndent=14*pt, spaceAfter=4*pt, leading=17)
                while i < len(lines):
                    bl = lines[i].strip()
                    if re.match(r'^[-*] ', bl):
                        story.append(Paragraph(f"• {inline(bl[2:])}", s_bullet))
                        i += 1
                    elif bl == '':
                        j = i + 1
                        while j < len(lines) and not lines[j].strip():
                            j += 1
                        if j < len(lines) and re.match(r'^[-*] ', lines[j].strip()):
                            i += 1
                        else:
                            break
                    else:
                        break
                _inject_chart_inline(_CHART_LABEL_RE.sub('', h3_text).strip())
            else:
                story.append(Paragraph(inline(h3_text), s_h3))
                i += 1

        elif re.match(r'^#{4,} ', stripped):
            h4_text = re.sub(r'^#{4,} ', '', stripped)
            if _CHART_LABEL_RE.match(h4_text):
                # 四级及以下的图表标签标题：解读+图表同块
                i = _handle_chart_label_block(h4_text, s_h3, i + 1)
            else:
                # 普通四级标题：渲染为 h3 样式
                story.append(Paragraph(inline(h4_text), s_h3))
                i += 1

        elif re.match(r'^-{3,}$', stripped):
            story.append(HRFlowable(width="100%", thickness=0.5, color=C_BDR))
            story.append(Spacer(1, 4*pt))
            i += 1

        elif re.match(r'^[-*] ', stripped):
            # ── 无序列表：合并软换行续行，再创建 bullet 段落
            while i < len(lines) and re.match(r'^[-*] ', lines[i].strip()):
                bullet_parts = [lines[i].strip()[2:]]
                i += 1
                while i < len(lines) and not _is_block_start(lines[i]):
                    bullet_parts.append(lines[i].strip())
                    i += 1
                story.append(Paragraph(f"• {inline(''.join(bullet_parts))}", s_bullet))

        elif re.match(r'^\d+\. ', stripped):
            # ── 有序列表：合并软换行续行
            num = 1
            while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
                num_parts = [re.sub(r'^\d+\. ', '', lines[i].strip())]
                i += 1
                while i < len(lines) and not _is_block_start(lines[i]):
                    num_parts.append(lines[i].strip())
                    i += 1
                story.append(Paragraph(f"{num}. {inline(''.join(num_parts))}", s_num))
                num += 1

        elif stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            # 跳过分隔行（|:---|:---:|）
            tbl_lines = [l for l in tbl_lines
                         if not re.match(r'^\|[\s\-|:]+\|$', l)]
            if tbl_lines:
                raw_data = []
                for tl in tbl_lines:
                    raw_data.append([c.strip() for c in tl.strip("|").split("|")])
                n_cols = max(len(r) for r in raw_data)
                for row in raw_data:
                    while len(row) < n_cols:
                        row.append("")

                # ── 自动列宽：按内容长度比例分配，避免换行 ────────────────────
                # 估算：CJK 字符宽 ≈ 字号 pt；ASCII/数字 ≈ 字号 × 0.55 pt；每列加 16pt 内边距
                _fs = 9  # 表格字号
                def _est_w(text: str) -> float:
                    cjk = sum(1 for ch in text if '一' <= ch <= '鿿' or '　' <= ch <= '〿')
                    asc = len(text) - cjk
                    return cjk * _fs + asc * _fs * 0.55 + 16  # +16 为左右内边距

                max_ws = [0.0] * n_cols
                for row in raw_data:
                    for ci, cell in enumerate(row):
                        # 先格式化再量宽（千位符会增加字符）
                        max_ws[ci] = max(max_ws[ci], _est_w(_add_thousands(cell)))

                total_est = sum(max_ws) or 1
                # 按比例分配，最小列宽 14mm，总宽不超 avail_w
                _14mm = 14 * mm
                col_widths = [max(w / total_est * avail_w, _14mm) for w in max_ws]
                # 如果等比超出，等比缩放
                if sum(col_widths) > avail_w:
                    scale = avail_w / sum(col_widths)
                    col_widths = [w * scale for w in col_widths]

                pdf_data = []
                for ri, row in enumerate(raw_data):
                    sty = s_hcell if ri == 0 else s_cell
                    pdf_data.append([Paragraph(inline(c), sty) for c in row])
                t = Table(pdf_data, colWidths=col_widths, repeatRows=1)
                ts = [
                    # 表头：深海蓝背景 + 白色文字（对齐 HTML）
                    ("BACKGROUND",    (0, 0),  (-1, 0),  C_NAVY),
                    ("TEXTCOLOR",     (0, 0),  (-1, 0),  white),
                    ("FONTNAME",      (0, 0),  (-1, -1), FONT),
                    ("FONTSIZE",      (0, 0),  (-1, -1), 9),
                    ("BOLD",          (0, 0),  (-1, 0),  True),
                    # 边框：仅底部分割线，无竖线（接近 HTML 表格风格）
                    ("LINEBELOW",     (0, 0),  (-1, -2), 0.5, C_BDR),
                    ("TOPPADDING",    (0, 0),  (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0),  (-1, -1), 6),
                    ("LEFTPADDING",   (0, 0),  (-1, -1), 8),
                    ("RIGHTPADDING",  (0, 0),  (-1, -1), 8),
                ]
                for ri in range(1, len(pdf_data)):
                    bg = C_LIGHT if ri % 2 == 0 else white
                    ts.append(("BACKGROUND", (0, ri), (-1, ri), bg))
                t.setStyle(TableStyle(ts))
                story.append(t)
                story.append(Spacer(1, 8*pt))

        elif _CHART_LABEL_RE.match(stripped):
            # ── 图表标签行（如 "图表1:xxx"）: 解读文字 + 图表打包为 KeepTogether 块
            i = _handle_chart_label_block(stripped, s_body, i + 1)

        elif _MD_IMG_RE.match(stripped):
            # Markdown 图片占位符 ![图表N:caption]() → 注入匹配图表（不输出文字）
            raw_cap = _MD_IMG_RE.match(stripped).group(1)
            caption_hint = re.sub(r'^图表?\s*\d+\s*[：:]', '', raw_cap).strip()
            _inject_chart_inline(caption_hint)
            i += 1

        elif stripped == "":
            story.append(Spacer(1, 3*pt))
            i += 1

        else:
            # 合并连续正文行：Markdown 软换行（单 \n）在段落内不产生段落分隔
            para_parts = [stripped]
            i += 1
            while i < len(lines) and not _is_block_start(lines[i]):
                para_parts.append(lines[i].strip())
                i += 1
            story.append(Paragraph(inline(''.join(para_parts)), s_body))

    # 最后一个章节的图表
    _add_h3_chart()      # 处理最后一个 h3 子节的图表
    _add_sec_chart()     # 处理最后一个 ## 章节的图表（兜底）

    # 未匹配的图表统一追加到附录
    _unplaced = [c for c in _pdf_charts if not c['placed']]
    if _unplaced:
        story.append(Spacer(1, 12 * pt))
        story.append(HRFlowable(width="100%", thickness=1, color=C_BDR))
        story.append(Spacer(1, 4 * pt))
        story.append(Paragraph("附图", s_h1))
        story.append(HRFlowable(width="100%", thickness=0.5, color=C_BDR, spaceAfter=4 * pt))
        for _c in _unplaced:
            _png = _c.get('png')
            if _png is not None:
                from reportlab.platypus import Image as _RLImage
                story.append(_RLImage(BytesIO(_png), width=avail_w, height=avail_w * 440 / 960))
                story.append(Paragraph(f"图: {inline(_c['caption'])}", s_date))
                story.append(Spacer(1, 8 * pt))

    doc.build(story)
    return buf.getvalue()


def html_to_pdf(html_content: str) -> bytes:
    """保留供外部调用；内部已改用 markdown_to_pdf（reportlab）。"""
    raise NotImplementedError(
        "html_to_pdf 已废弃，请改用 markdown_to_pdf(markdown_content, title)。"
    )
